"""
Модуль анализа документов через OpenRouter API.
Vision-модель для сканов/изображений, текстовая модель для текстовых документов.
Батчинг через asyncio.
"""

import asyncio
import base64
import hashlib
import json
import io
import logging
import shutil
import subprocess
import tempfile

from doctypes import get_prompt_doc_types
from pathlib import Path

import httpx
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from openpyxl import load_workbook

from scanner import get_file_type

OPENROUTER_URL = "https://openrouter.ai/api/v1/chat/completions"

_DOC_TYPE_LIST = get_prompt_doc_types()

ANALYSIS_PROMPT = f"""Ты — помощник юриста, анализирующий документы. Проанализируй предоставленный документ и верни СТРОГО JSON (без markdown, без ```json):

{{
  "doc_type": "тип документа ({_DOC_TYPE_LIST})",
  "title": "полное название документа",
  "number": "номер документа или пустая строка",
  "date": "дата документа в формате ДД.ММ.ГГГГ или пустая строка",
  "party_1_name": "название/ФИО стороны 1 (составитель/инициатор документа: Продавец, Исполнитель, Подрядчик, Поставщик)",
  "party_1_role": "роль стороны 1 (Продавец / Исполнитель / Подрядчик / Поставщик / Доверитель / и т.д.) или пустая строка",
  "party_2_name": "название/ФИО стороны 2 (контрагент: Покупатель, Заказчик, Клиент)",
  "party_2_role": "роль стороны 2 (Покупатель / Заказчик / Клиент / Поверенный / и т.д.) или пустая строка",
  "reference_number": "номер ДОГОВОРА или иного базового документа, к которому относится данный документ (например, для УПД — номер договора поставки). НЕ указывайте собственный номер документа. Или пустая строка",
  "reference_date": "дата базового документа в формате ДД.ММ.ГГГГ или пустая строка",
  "amount": "сумма документа цифрами (напр. '1 500 000,00 руб.') или пустая строка",
  "goods_summary": "краткое описание предмета документа (товары, услуги, работы) или пустая строка",
  "summary": "краткое содержание (1-2 предложения)",
  "is_multidoc": true/false
}}

Важно:
- Если какое-то поле не удаётся определить — оставь пустую строку
- Тип документа выбирай строго из списка выше
- Спецификация — это перечень товаров/услуг к договору, НЕ дополнительное соглашение
- График поставки — это расписание поставок по договору, НЕ акт
- Дату приводи к формату ДД.ММ.ГГГГ
- party_1 — сторона, составляющая/инициирующая документ; party_2 — контрагент
- reference_number — это номер ДОГОВОРА/БАЗОВОГО документа, на который ссылается данный документ (например, для УПД укажи номер договора). НЕ указывайте собственный номер документа в это поле
- amount — полная сумма как написано в документе, включая валюту. Не придумывай суммы если их нет в тексте
- goods_summary — перечисли основные товары/услуги (до 200 символов)
- is_multidoc=true если файл содержит НЕСКОЛЬКО разных документов подряд. Обычные многостраничные документы одного типа — false
- Верни ТОЛЬКО JSON, без пояснений
"""


_PDF_TEXT_MIN_LENGTH = 100  # Минимум символов, чтобы считать PDF текстовым
_TAIL_PAGES = 3  # Сколько страниц брать с конца для текстовых PDF


def _extract_pdf_text(pdf_path: Path, max_pages: int = 10) -> str:
    """Извлекает текст из PDF: первые max_pages страниц + последние _TAIL_PAGES.

    Для юридических документов реквизиты, подписи и суммы обычно в конце.
    """
    doc = fitz.open(str(pdf_path))
    total = doc.page_count
    texts: dict[int, str] = {}

    # Голова — первые max_pages страниц
    head_end = min(max_pages, total)
    for i in range(head_end):
        page_text = doc[i].get_text("text").strip()
        if page_text:
            texts[i] = page_text

    # Хвост — последние _TAIL_PAGES страниц (если не пересекаются с головой)
    tail_start = max(head_end, total - _TAIL_PAGES)
    for i in range(tail_start, total):
        if i not in texts:
            page_text = doc[i].get_text("text").strip()
            if page_text:
                texts[i] = page_text

    doc.close()
    # Склеиваем по порядку страниц
    return "\n\n".join(texts[k] for k in sorted(texts))


def _sample_page_indices(total_pages: int, max_pages: int) -> list[int]:
    """Выбирает оптимальные индексы страниц для анализа.

    Вместо первых N страниц берёт: первую, последнюю,
    и равномерно распределённые промежуточные.
    """
    if total_pages <= max_pages:
        return list(range(total_pages))
    if max_pages == 1:
        return [0]
    if max_pages == 2:
        return [0, total_pages - 1]
    indices = set()
    indices.add(0)               # первая
    indices.add(total_pages - 1) # последняя
    remaining = max_pages - len(indices)
    for k in range(1, remaining + 1):
        pos = round(total_pages * k / (remaining + 1))
        pos = max(1, min(pos, total_pages - 2))
        indices.add(pos)
    return sorted(indices)


def _pdf_to_images(pdf_path: Path, max_pages: int = 3) -> list[bytes]:
    """Рендерит выбранные страницы PDF в PNG-изображения (умный сэмплинг)."""
    doc = fitz.open(str(pdf_path))
    total = doc.page_count
    page_indices = _sample_page_indices(total, max_pages)
    images = []
    for i in page_indices:
        if i < total:
            pix = doc[i].get_pixmap(dpi=200)
            images.append(pix.tobytes("png"))
    doc.close()
    return images


def _pdf_page_count(pdf_path: Path) -> int:
    try:
        doc = fitz.open(str(pdf_path))
        n = doc.page_count
        doc.close()
        return n
    except Exception:
        return 0


def check_suspicious(doc: dict, page_threshold: int = 10) -> tuple[bool, str]:
    """
    Проверяет, является ли документ подозрительным (требующим внимания).
    Возвращает (is_suspicious, reason).
    """
    if doc.get("is_multidoc"):
        return True, "содержит несколько документов"

    if doc.get("_ext") == ".pdf" and doc.get("_page_count", 0) > page_threshold:
        return True, f"большой PDF: {doc['_page_count']} стр."

    if not (doc.get("title") or "").strip():
        return True, "название не определено"

    return False, ""


def _image_to_base64(image_bytes: bytes) -> str:
    return base64.b64encode(image_bytes).decode("utf-8")


def _read_image_file(path: Path) -> bytes:
    with open(path, "rb") as f:
        return f.read()


def _docx_page_count(path: Path) -> int:
    """Оценка количества страниц DOCX по параграфам."""
    try:
        doc = DocxDocument(str(path))
        paragraphs = len([p for p in doc.paragraphs if p.text.strip()])
        return max(1, paragraphs // 25)
    except Exception:
        return 1


def _extract_docx_text(path: Path) -> str:
    """Извлекает текст из DOCX. Для старого .doc пытается прочитать как binary."""
    try:
        doc = DocxDocument(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs[:100])
    except Exception:
        # .doc (старый формат) — python-docx его не читает, пробуем как binary
        return _extract_binary_text(path)


def _extract_binary_text(path: Path) -> str:
    """Извлекает текст из бинарных форматов (.doc, .rtf) — грубое извлечение."""
    try:
        with open(path, "rb") as f:
            raw = f.read(100000)
        # Пытаемся декодировать как UTF-8/CP1251, фильтруем printable
        for enc in ("utf-8", "cp1251", "latin-1"):
            try:
                text = raw.decode(enc, errors="ignore")
                # Убираем непечатные символы, оставляем читаемый текст
                import re
                text = re.sub(r"[^\S\n\r]+", " ", text)  # множественные пробелы → один
                lines = [l.strip() for l in text.splitlines() if l.strip() and len(l.strip()) > 3]
                return "\n".join(lines[:150])
            except Exception:
                continue
    except Exception:
        pass
    return ""


def _convert_doc_to_pdf(path: Path) -> Path | None:
    """Конвертирует .doc/.docx в PDF через LibreOffice. Возвращает путь к PDF или None."""
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            result = subprocess.run(
                [
                    "soffice", "--headless", "--convert-to", "pdf",
                    "--outdir", tmpdir,
                    str(path),
                ],
                capture_output=True, timeout=60,
            )
            if result.returncode != 0:
                logging.warning("soffice convert failed for %s: %s", path, result.stderr.decode(errors="ignore"))
                return None
            # Ищем созданный PDF
            pdf_name = Path(path).stem + ".pdf"
            pdf_path = Path(tmpdir) / pdf_name
            if pdf_path.exists():
                # Копируем в уникальный временный файл, чтобы не удалить при выходе из tempdir
                unique = hashlib.md5(str(path).encode()).hexdigest()[:8]
                persist = Path(tempfile.gettempdir()) / f"docsorter_{unique}_{pdf_name}"
                shutil.copy2(pdf_path, persist)
                return persist
            return None
    except Exception as e:
        logging.warning("soffice convert error for %s: %s", path, e)
        return None


def _extract_rtf_text(path: Path) -> str:
    """Извлекает текст из RTF — убирает разметку."""
    try:
        with open(path, "r", encoding="cp1251", errors="ignore") as f:
            raw = f.read(50000)
    except Exception:
        try:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                raw = f.read(50000)
        except Exception:
            return ""
    # Убираем RTF-разметку
    import re
    # Убираем управляющие слова {\word}, \viewkind4, etc.
    text = re.sub(r'\\[a-z]+\d*\s?', ' ', raw)
    # Убираем фигурные скобки
    text = text.replace('{', ' ').replace('}', ' ')
    # Убираем лишние пробелы
    text = re.sub(r'\s+', ' ', text)
    return text.strip()[:10000]


def _extract_xlsx_text(path: Path) -> str:
    wb = load_workbook(str(path), read_only=True, data_only=True)
    lines = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        lines.append(f"=== Лист: {sheet_name} ===")
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i >= 50:  # Первые 50 строк
                break
            cells = [str(c) if c is not None else "" for c in row]
            lines.append(" | ".join(cells))
    wb.close()
    return "\n".join(lines)


def _extract_text(path: Path) -> str:
    try:
        with open(path, "r", encoding="utf-8") as f:
            return f.read(10000)
    except UnicodeDecodeError:
        with open(path, "r", encoding="cp1251") as f:
            return f.read(10000)


def _get_mime_type(ext: str) -> str:
    mapping = {
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".png": "image/png",
        ".tiff": "image/tiff",
        ".tif": "image/tiff",
        ".bmp": "image/bmp",
        ".webp": "image/webp",
    }
    return mapping.get(ext.lower(), "image/png")


def _build_vision_messages(images: list[bytes], ext: str = ".png") -> list[dict]:
    """Формирует сообщения для vision-модели."""
    content = [{"type": "text", "text": ANALYSIS_PROMPT}]
    mime = _get_mime_type(ext)
    for img in images:
        b64 = _image_to_base64(img)
        content.append({
            "type": "image_url",
            "image_url": {
                "url": f"data:{mime};base64,{b64}",
            },
        })
    return [{"role": "user", "content": content}]


def _extract_json(raw_text: str) -> dict:
    """Извлекает JSON из ответа LLM с несколькими уровнями fallback."""
    import re as _re

    # Шаг 1: Прямой парс
    try:
        return json.loads(raw_text)
    except (json.JSONDecodeError, ValueError):
        pass

    # Шаг 2: Убираем markdown-обёртку
    if raw_text.startswith("```"):
        lines = raw_text.split("\n")
        lines = [l for l in lines if not l.strip().startswith("```")]
        cleaned = "\n".join(lines)
        try:
            return json.loads(cleaned)
        except (json.JSONDecodeError, ValueError):
            pass

    # Шаг 3: Regex — ищем JSON-объект
    match = _re.search(r'\{[\s\S]*\}', raw_text)
    if match:
        try:
            return json.loads(match.group(0))
        except (json.JSONDecodeError, ValueError):
            pass

    # Шаг 4: Regex — ищем JSON-массив
    match = _re.search(r'\[[\s\S]*\]', raw_text)
    if match:
        try:
            return json.loads(match.group(0))
        except (json.JSONDecodeError, ValueError):
            pass

    raise json.JSONDecodeError("Не удалось извлечь JSON из ответа LLM", raw_text, 0)


def _build_text_messages(text: str) -> list[dict]:
    """Формирует сообщения для текстовой модели."""
    return [
        {
            "role": "user",
            "content": f"{ANALYSIS_PROMPT}\n\nТекст документа:\n\n{text}",
        }
    ]


async def _call_openrouter(
    client: httpx.AsyncClient,
    api_key: str,
    model: str,
    messages: list[dict],
    json_mode: bool = True,
) -> tuple[dict, httpx.Response]:
    """Вызов OpenRouter API. Возвращает (parsed_json, raw_response)."""
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://docsorter.app",
        "X-Title": "DocSorter",
    }
    payload: dict = {
        "model": model,
        "messages": messages,
        "temperature": 0.1,
        "max_tokens": 1000,
    }
    if json_mode:
        payload["response_format"] = {"type": "json_object"}
    resp = await client.post(
        OPENROUTER_URL,
        json=payload,
        headers=headers,
        timeout=120.0,
    )
    resp.raise_for_status()
    data = resp.json()

    raw_text = data["choices"][0]["message"]["content"].strip()
    return _extract_json(raw_text), resp


async def _call_with_retry(
    client: httpx.AsyncClient,
    api_key: str,
    model: str,
    messages: list[dict],
    fallback_model: str = "",
    max_retries: int = 4,
) -> dict:
    """Вызов OpenRouter API с ретраями, Retry-After и fallback на другую модель."""
    last_exc = None
    for attempt in range(max_retries + 1):
        try:
            result, _resp = await _call_openrouter(client, api_key, model, messages)
            return result
        except httpx.HTTPStatusError as e:
            last_exc = e
            if attempt < max_retries:
                # Уважаем Retry-After от сервера
                retry_after = e.response.headers.get("retry-after")
                if retry_after:
                    try:
                        wait = float(retry_after)
                    except ValueError:
                        wait = 5.0
                else:
                    wait = min(2 ** attempt, 30)  # 1, 2, 4, 8 сек, макс 30
                logging.warning(
                    "HTTP %d, retry %d/%d for %s (wait %.0fs)",
                    e.response.status_code, attempt + 1, max_retries, model, wait,
                )
                await asyncio.sleep(wait)
        except (json.JSONDecodeError, Exception) as e:
            last_exc = e
            if attempt < max_retries:
                wait = min(2 ** attempt, 30)
                logging.warning(
                    "Retry %d/%d for model %s: %s", attempt + 1, max_retries, model, e
                )
                await asyncio.sleep(wait)

    # Fallback на другую модель
    if fallback_model and fallback_model != model:
        logging.warning("Falling back to model %s", fallback_model)
        try:
            result, _resp = await _call_openrouter(client, api_key, fallback_model, messages)
            return result
        except Exception as e:
            last_exc = e

    raise last_exc


def _normalize_party_name(name: str) -> str:
    """Заменяет полные организационно-правовые формы на аббревиатуры."""
    if not name:
        return name
    import re as _re
    # Порядок важен: более длинные сначала (ОАО, ПАО, ЗАО перед АО)
    _REPLACEMENTS = [
        # Организационно-правовые формы
        (r'Публичное\s+акционерное\s+общество', 'ПАО'),
        (r'Открытое\s+акционерное\s+общество', 'ОАО'),
        (r'Закрытое\s+акционерное\s+общество', 'ЗАО'),
        (r'Акционерное\s+общество', 'АО'),
        (r'Общество\s+с\s+ограниченной\s+ответственностью', 'ООО'),
        (r'Общество\s+с\s+дополнительной\s+ответственностью', 'ОДО'),
        (r'Индивидуальный\s+предприниматель', 'ИП'),
        (r'Некоммерческая\s+организация', 'НКО'),
        (r'Государственное\s+унитарное\s+предприятие', 'ГУП'),
        (r'Муниципальное\s+унитарное\s+предприятие', 'МУП'),
        (r'Федеральное\s+государственное\s+унитарное\s+предприятие', 'ФГУП'),
        # Государственные органы
        (r'Инспекция\s+федеральной\s+налоговой\s+службы', 'ИФНС'),
        (r'Федеральная\s+налоговая\s+служба', 'ФНС'),
        (r'Федеральная\s+служба\s+судебных\s+приставов', 'ФССП'),
        (r'Федеральная\s+служба\s+по\s+финансовым\s+рынкам', 'ФСФР'),
        (r'Федеральная\s+служба\s+по\s+финансовому\s+мониторингу', 'Росфинмониторинг'),
        (r'Федеральная\s+антимонопольная\s+служба', 'ФАС'),
        (r'Федеральная\s+служба\s+по\s+экологическому[\s,]+технологическому\s+и\s+атомному\s+надзору', 'Ростехнадзор'),
        (r'Федеральная\s+служба\s+по\s+надзору\s+в\s+сфере\s+защиты\s+прав\s+потребителей\s+и\s+благополучия\s+человека', 'Роспотребнадзор'),
        (r'Федеральная\s+служба\s+государственной\s+регистрации[\s,]+кадастра\s+и\s+картографии', 'Росреестр'),
        (r'Федеральная\s+служба\s+по\s+труду\s+и\s+занятости', 'Роструд'),
        (r'Федеральная\s+таможенная\s+служба', 'ФТС'),
        (r'Министерство\s+финансов\s+Российской\s+Федерации', 'Минфин России'),
        (r'Министерство\s+экономического\s+развития\s+Российской\s+Федерации', 'Минэкономразвития России'),
        (r'Министерство\s+юстиции\s+Российской\s+Федерации', 'Минюст России'),
        # Суды
        (r'Федеральный\s+арбитражный\s+суд\s+([^\s,]+(?:\s+[^\s,]+){0,3})\s+округа', r'ФАС \1 округа'),
        (r'Арбитражный\s+суд\s+([^\s,]+(?:\s+[^\s,]+){0,3})\s+округа', r'АС \1 округа'),
        (r'Арбитражный\s+суд\s+г[\.\s]+Москвы', 'АС г. Москвы'),
        (r'Арбитражный\s+суд\s+г[\.\s]+Санкт-Петербург[а-я]*\s+и\s+Ленинградской\s+области', 'АС г. Санкт-Петербурга и ЛО'),
        (r'Арбитражный\s+суд\s+([^\s,]+(?:\s+[^\s,]+){0,2})\s+области', r'АС \1 области'),
        (r'Арбитражный\s+суд\s+([^\s,]+(?:\s+[^\s,]+){0,2})\s+края', r'АС \1 края'),
        (r'Арбитражный\s+суд\s+Республики\s+([^\s,]+(?:\s+[^\s,]+){0,2})', r'АС Республики \1'),
        (r'Арбитражный\s+суд', 'АС'),
        (r'Верховный\s+суд\s+Российской\s+Федерации', 'ВС РФ'),
        (r'Верховный\s+суд\s+РФ', 'ВС РФ'),
        (r'Верховный\s+суд', 'ВС'),
        (r'Конституционный\s+суд\s+Российской\s+Федерации', 'КС РФ'),
        (r'Конституционный\s+суд\s+РФ', 'КС РФ'),
        (r'Суд\s+общей\s+юрисдикции\s+([^\s,]+(?:\s+[^\s,]+){0,3})', r'СОЮ \1'),
        (r'Районный\s+суд\s+г[\.\s]+([^\s,]+(?:\s+[^\s,]+){0,2})', r'Райсуд г. \1'),
        (r'Мировой\s+судь[яьи]\s+судебного\s+участка', 'Мировой судья уч.'),
        # Прочие
        (r'Федеральное\s+казначейство', 'Казначейство России'),
        (r'Прокуратура\s+Российской\s+Федерации', 'Генпрокуратура'),
        (r'Прокуратура\s+РФ', 'Генпрокуратура'),
        (r'Следственный\s+комитет\s+Российской\s+Федерации', 'СК РФ'),
        (r'Следственный\s+комитет\s+РФ', 'СК РФ'),
    ]
    for pattern, abbr in _REPLACEMENTS:
        name = _re.sub(pattern, abbr, name, flags=_re.IGNORECASE)
    # Убираем лишние пробелы
    name = _re.sub(r'\s{2,}', ' ', name).strip()
    return name


def _postprocess_result(result: dict) -> dict:
    """Сериализует party-поля и синтезирует legacy-поля из новых структурированных."""
    # Сериализуем party_1_name + party_1_role → party_1 (JSON-строка)
    for pfield in ("party_1", "party_2"):
        name_key = f"{pfield}_name"
        role_key = f"{pfield}_role"
        name_val = result.pop(name_key, "")
        role_val = result.pop(role_key, "")
        # Нормализуем ОПФ → аббревиатура
        name_val = _normalize_party_name(name_val)
        if name_val:
            result[pfield] = json.dumps(
                {"name": name_val, "role": role_val}, ensure_ascii=False,
            )
        else:
            result[pfield] = ""

    # Синтезируем legacy counterparty
    if not result.get("counterparty"):
        p1 = _parse_party_name(result.get("party_1", ""))
        p2 = _parse_party_name(result.get("party_2", ""))
        result["counterparty"] = p1 or p2 or ""

    # Синтезируем legacy reference
    if not result.get("reference"):
        ref_num = result.get("reference_number", "").strip()
        ref_date = result.get("reference_date", "").strip()
        if ref_num:
            ref = f"к документу №{ref_num}"
            if ref_date:
                ref += f" от {ref_date}"
            result["reference"] = ref
        else:
            result["reference"] = ""

    return result


def _parse_party_name(raw: str) -> str:
    """Извлекает имя из JSON-строки party-поля."""
    if not raw:
        return ""
    try:
        data = json.loads(raw)
        return data.get("name", "")
    except (json.JSONDecodeError, TypeError):
        return raw


async def analyze_file(
    file_info: dict,
    api_key: str,
    vision_model: str,
    text_model: str,
    max_pages: int,
    semaphore: asyncio.Semaphore,
    client: httpx.AsyncClient,
    error_callback=None,
) -> dict:
    """Анализирует один файл и возвращает метаданные."""
    async with semaphore:
        path = file_info["path"]
        ext = file_info["ext"]
        file_type = get_file_type(ext)

        try:
            if file_type == "image":
                img_data = _read_image_file(path)
                messages = _build_vision_messages([img_data], ext)
                result = await _call_with_retry(
                    client, api_key, vision_model, messages, fallback_model=text_model,
                )

            elif file_type == "pdf":
                # Сначала пробуем извлечь текст через PyMuPDF
                text = _extract_pdf_text(path, max_pages)
                if len(text.strip()) >= _PDF_TEXT_MIN_LENGTH:
                    # Текстовый PDF — отправляем в text_model (точнее и дешевле)
                    messages = _build_text_messages(text)
                    result = await _call_with_retry(
                        client, api_key, text_model, messages, fallback_model=vision_model,
                    )
                else:
                    # Скан/картинка — fallback на VL-модель
                    images = _pdf_to_images(path, max_pages)
                    if images:
                        messages = _build_vision_messages(images, ".png")
                        result = await _call_with_retry(
                            client, api_key, vision_model, messages, fallback_model=text_model,
                        )
                    else:
                        result = _empty_result("Пустой PDF")

            elif file_type == "docx":
                # Для старых .doc (OLE2) — конвертация в PDF через LibreOffice
                if ext.lower() == ".doc":
                    pdf_path = _convert_doc_to_pdf(Path(path))
                    if pdf_path:
                        try:
                            text = _extract_pdf_text(pdf_path, max_pages)
                            if len(text.strip()) >= _PDF_TEXT_MIN_LENGTH:
                                messages = _build_text_messages(text)
                                result = await _call_with_retry(
                                    client, api_key, text_model, messages, fallback_model=vision_model,
                                )
                            else:
                                images = _pdf_to_images(pdf_path, max_pages)
                                if images:
                                    messages = _build_vision_messages(images, ".png")
                                    result = await _call_with_retry(
                                        client, api_key, vision_model, messages, fallback_model=text_model,
                                    )
                                else:
                                    result = _empty_result("Пустой документ (конвертация)")
                        finally:
                            # Удаляем временный PDF
                            try:
                                pdf_path.unlink(missing_ok=True)
                            except Exception:
                                pass
                    else:
                        # Fallback: пробуем извлечь текст как binary (старый метод)
                        text = _extract_binary_text(Path(path))
                        if text.strip():
                            messages = _build_text_messages(text)
                            result = await _call_with_retry(
                                client, api_key, text_model, messages, fallback_model=vision_model,
                            )
                        else:
                            result = _empty_result("Не удалось прочитать .doc")
                else:
                    # .docx — читаем через python-docx
                    text = _extract_docx_text(path)
                    if text.strip():
                        messages = _build_text_messages(text)
                        result = await _call_with_retry(
                            client, api_key, text_model, messages, fallback_model=vision_model,
                        )
                    else:
                        result = _empty_result("Пустой документ")

            elif file_type == "xlsx":
                text = _extract_xlsx_text(path)
                if text.strip():
                    messages = _build_text_messages(text)
                    result = await _call_with_retry(
                        client, api_key, text_model, messages, fallback_model=vision_model,
                    )
                else:
                    result = _empty_result("Пустая таблица")

            elif file_type == "rtf":
                text = _extract_rtf_text(path)
                if text.strip():
                    messages = _build_text_messages(text)
                    result = await _call_with_retry(
                        client, api_key, text_model, messages, fallback_model=vision_model,
                    )
                else:
                    result = _empty_result("Пустой RTF")

            elif file_type == "text":
                text = _extract_text(path)
                if text.strip():
                    messages = _build_text_messages(text)
                    result = await _call_with_retry(
                        client, api_key, text_model, messages, fallback_model=vision_model,
                    )
                else:
                    result = _empty_result("Пустой файл")

            else:
                result = _empty_result("Неподдерживаемый формат")

            # Постобработка: сериализация party-полей и синтез legacy-полей
            result = _postprocess_result(result)

        except (json.JSONDecodeError, httpx.HTTPStatusError) as e:
            reason = "Ошибка разбора ответа LLM"
            if isinstance(e, httpx.HTTPStatusError):
                reason = f"Ошибка API: {e.response.status_code}"
            result = _empty_result(reason)
            logging.warning("Ошибка анализа %s: %s", file_info.get("name", "?"), reason)
            if error_callback:
                error_callback(f"⚠ {file_info.get('name', '?')}: {reason}")
        except Exception as e:
            reason = f"Ошибка: {str(e)[:100]}"
            result = _empty_result(reason)
            logging.warning("Ошибка анализа %s: %s", file_info.get("name", "?"), reason)
            if error_callback:
                error_callback(f"⚠ {file_info.get('name', '?')}: {reason}")

        # Дополняем информацией о файле
        result["_file_name"] = file_info["name"]
        result["_file_path"] = str(file_info["path"])
        result["_rel_path"] = file_info["rel_path"]
        result["_ext"] = ext
        result["_file_hash"] = file_info.get("hash", "")
        result["_page_count"] = file_info.get("page_count", 0)

        # Подозрительность
        threshold = file_info.get("suspicious_threshold", 10)
        suspicious, reason = check_suspicious(result, threshold)
        result["_suspicious"] = suspicious
        result["_suspicious_reason"] = reason

        return result


def _empty_result(reason: str) -> dict:
    return {
        "doc_type": "Прочее",
        "title": reason,
        "number": "",
        "date": "",
        "counterparty": "",
        "reference": "",
        "reference_number": "",
        "reference_date": "",
        "party_1": "",
        "party_2": "",
        "amount": "",
        "goods_summary": "",
        "summary": reason,
        "is_multidoc": False,
    }


async def analyze_batch(
    files: list[dict],
    api_key: str,
    vision_model: str,
    text_model: str,
    max_pages: int = 3,
    max_concurrent: int = 5,
    suspicious_threshold: int = 10,
    progress_callback=None,
    error_callback=None,
) -> list[dict]:
    """
    Анализирует все файлы с батчингом.
    progress_callback(current, total) вызывается после каждого файла.
    error_callback(message) вызывается при ошибке анализа файла.
    """
    from project import file_hash

    semaphore = asyncio.Semaphore(max_concurrent)
    results = []

    # Предварительно считаем хэши и количество страниц
    for f in files:
        f["hash"] = file_hash(f["path"])
        f["suspicious_threshold"] = suspicious_threshold
        if f["ext"] == ".pdf":
            f["page_count"] = _pdf_page_count(f["path"])
        elif f["ext"] in (".docx", ".doc"):
            f["page_count"] = _docx_page_count(f["path"])
        else:
            f["page_count"] = 1

    # Дедупликация по хэшу: анализируем только уникальные файлы
    unique_files = []
    hash_to_unique_idx = {}  # hash → индекс в unique_files
    duplicate_map = {}  # индекс в оригинальном files → индекс в unique_files
    for i, f in enumerate(files):
        h = f.get("hash", "")
        if h and h in hash_to_unique_idx:
            duplicate_map[i] = hash_to_unique_idx[h]
        else:
            if h:
                hash_to_unique_idx[h] = len(unique_files)
            unique_files.append(f)

    async with httpx.AsyncClient() as client:
        tasks = []
        for file_info in unique_files:
            task = analyze_file(
                file_info, api_key, vision_model, text_model,
                max_pages, semaphore, client, error_callback,
            )
            tasks.append(task)

        # Собираем результаты уникальных файлов
        unique_results = []
        completed = 0
        for coro in asyncio.as_completed(tasks):
            result = await coro
            unique_results.append(result)
            completed += 1
            if progress_callback:
                progress_callback(completed, len(unique_files))

    # Маппинг: путь → результат (для уникальных)
    path_to_result = {str(uf["path"]): r for uf, r in zip(unique_files, unique_results)}

    # Собираем финальный список: для дубликатов клонируем результат оригинала
    import copy
    for i, f in enumerate(files):
        if i in duplicate_map:
            orig_path = str(unique_files[duplicate_map[i]]["path"])
            result = copy.deepcopy(path_to_result.get(orig_path, _empty_result("Дубликат")))
            # Подменяем пути на дубликат
            result["_file_name"] = f["name"]
            result["_file_path"] = str(f["path"])
            result["_rel_path"] = f["rel_path"]
            result["_ext"] = f["ext"]
            result["_file_hash"] = f.get("hash", "")
            result["_page_count"] = f.get("page_count", 0)
        else:
            result = path_to_result.get(str(f["path"]), _empty_result("Не найден"))
        results.append(result)

    # Сортируем по оригинальному порядку файлов
    file_order = {str(f["path"]): i for i, f in enumerate(files)}
    results.sort(key=lambda r: file_order.get(r.get("_file_path", ""), 0))

    return results
